from db_collection import db_collection
from docx import Document
from docx.shared import Inches

class learner(object):
	def __init__(self):
		self.document = Document("cmpe273-greensheet.docx")
		self.i = 0
		self.insert = db_collection()
		
	def read(self):
		for table in self.document.tables:
			self.i = 0
			while(self.i < len(table.rows)):
				self.insert.insertCollection("%s for %s is %s" % (table.cell(self.i,0).text.encode('utf-8'), table.cell(self.i,1).text.encode('utf-8'), table.cell(self.i,2).text.encode('utf-8')))
				self.i = self.i + 1
		
		for paragraph in self.document.paragraphs:
			if (paragraph.text.startswith(":")):
				self.insert.insertCollection(paragraph.text.encode('utf-8').replace(":", ""))
			elif (" ** " in paragraph.text):
				self.insert.insertCollection(paragraph.text.encode('utf-8').replace("**", "is"))
			else :
				self.insert.insertCollection(paragraph.text.encode('utf-8'))